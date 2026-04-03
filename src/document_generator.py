"""
document_generator.py
=====================
MoM text ko PDF, Word (.docx), aur Markdown format mein export karta hai.
"""

import os
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Emoji characters that standard PDF fonts cannot render
_EMOJI_RE = re.compile(
    "[\U00010000-\U0010ffff"   # Supplementary multilingual plane (most emojis)
    "\U0001f300-\U0001f9ff"    # Misc symbols and pictographs
    "\U00002600-\U000027bf"    # Misc symbols
    "\U0000fe00-\U0000fe0f"    # Variation selectors
    "\U00020000-\U0002a6df"    # CJK extension B
    "]",
    flags=re.UNICODE,
)

# Map emoji section markers to text equivalents for PDF output
_SECTION_EMOJI_MAP = {
    "📅": "[DATE]",
    "⏱️": "[DURATION]",
    "👥": "[PARTICIPANTS]",
    "📌": "[KEY POINTS]",
    "✅": "[ACTION ITEMS]",
    "📝": "[DECISIONS]",
    "⚠️": "[OPEN ISSUES]",
    "📆": "[NEXT STEPS]",
}


def _sanitize_for_pdf(text: str) -> str:
    """
    PDF ke liye text sanitize karta hai.
    Emojis ko text equivalents se replace karta hai aur
    latin-1 mein encode nahi hone wale characters ko safely handle karta hai.
    """
    # Replace known section emojis with readable text
    for emoji, replacement in _SECTION_EMOJI_MAP.items():
        text = text.replace(emoji, replacement)
    # Remove any remaining emoji/supplementary characters
    text = _EMOJI_RE.sub("", text)
    # Encode to latin-1 safely (replaces unknown chars with ?)
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text


class DocumentGenerator:
    """
    MoM text se professional documents generate karta hai.

    Supported formats:
        - PDF  (via fpdf2)
        - DOCX (via python-docx)
        - Markdown (.md)

    Usage::

        gen = DocumentGenerator(output_dir="outputs")
        pdf_path  = gen.to_pdf(mom_text)
        docx_path = gen.to_docx(mom_text)
        md_path   = gen.to_markdown(mom_text)
    """

    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def to_pdf(
        self,
        mom_text: str,
        filename: Optional[str] = None,
        title: str = "Minutes of Meeting",
    ) -> str:
        """
        MoM text ko PDF file mein save karta hai.

        Args:
            mom_text: Formatted MoM text.
            filename: Output filename (without extension). Auto-generated if None.
            title: PDF document title.

        Returns:
            Path to generated PDF file.
        """
        try:
            from fpdf import FPDF
        except ImportError as exc:
            raise ImportError(
                "fpdf2 install nahi hai! Run: pip install fpdf2"
            ) from exc

        output_path = self._get_output_path(filename, ".pdf")

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        # Title
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 12, title, ln=True, align="C")
        pdf.ln(4)

        # Timestamp
        pdf.set_font("Helvetica", "I", 10)
        pdf.cell(
            0,
            8,
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            ln=True,
            align="C",
        )
        pdf.ln(6)

        # Divider
        pdf.set_draw_color(100, 100, 100)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(6)

        # Content - process line by line
        for line in mom_text.split("\n"):
            # Strip markdown bold markers
            clean = line.replace("**", "").replace("__", "")

            if not clean.strip():
                pdf.ln(3)
                continue

            # Section headers (lines starting with emoji + bold)
            if any(
                clean.startswith(emoji)
                for emoji in ["📅", "⏱️", "👥", "📌", "✅", "📝", "⚠️", "📆"]
            ):
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(31, 73, 125)  # Dark blue
                pdf.multi_cell(0, 8, _sanitize_for_pdf(clean))
                pdf.set_text_color(0, 0, 0)
                pdf.ln(2)
            elif clean.startswith("---"):
                pdf.set_draw_color(200, 200, 200)
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica", "", 11)
                pdf.multi_cell(0, 7, _sanitize_for_pdf(clean))

        pdf.output(output_path)
        logger.info("PDF saved: %s", output_path)
        return output_path

    def to_docx(
        self,
        mom_text: str,
        filename: Optional[str] = None,
        title: str = "Minutes of Meeting",
    ) -> str:
        """
        MoM text ko Word (.docx) file mein save karta hai.

        Args:
            mom_text: Formatted MoM text.
            filename: Output filename (without extension). Auto-generated if None.
            title: Document title.

        Returns:
            Path to generated DOCX file.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise ImportError(
                "python-docx install nahi hai! Run: pip install python-docx"
            ) from exc

        output_path = self._get_output_path(filename, ".docx")
        doc = Document()

        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Timestamp
        ts_para = doc.add_paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts_para.runs[0].font.italic = True
        ts_para.runs[0].font.size = Pt(10)
        ts_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Spacer

        # Content
        for line in mom_text.split("\n"):
            clean = line.replace("**", "").replace("__", "")

            if not clean.strip():
                doc.add_paragraph()
                continue

            # Section headers
            if any(
                clean.startswith(emoji)
                for emoji in ["📅", "⏱️", "👥", "📌", "✅", "📝", "⚠️", "📆"]
            ):
                heading = doc.add_heading(clean, level=2)
                heading.runs[0].font.color.rgb = RGBColor(31, 73, 125)
            elif clean.startswith("---"):
                doc.add_paragraph("─" * 60)
            else:
                para = doc.add_paragraph(clean)
                para.runs[0].font.size = Pt(11) if para.runs else Pt(11)

        doc.save(output_path)
        logger.info("DOCX saved: %s", output_path)
        return output_path

    def to_markdown(
        self,
        mom_text: str,
        filename: Optional[str] = None,
        title: str = "Minutes of Meeting",
    ) -> str:
        """
        MoM text ko Markdown (.md) file mein save karta hai.

        Args:
            mom_text: Formatted MoM text.
            filename: Output filename (without extension). Auto-generated if None.
            title: Document title.

        Returns:
            Path to generated Markdown file.
        """
        output_path = self._get_output_path(filename, ".md")

        header = (
            f"# {title}\n\n"
            f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
            "---\n\n"
        )

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(header + mom_text)

        logger.info("Markdown saved: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _get_output_path(self, filename: Optional[str], extension: str) -> str:
        """Output file ka full path generate karta hai."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"MoM_{timestamp}"
        # Remove extension if already included
        filename = Path(filename).stem
        return str(Path(self.output_dir) / (filename + extension))
